import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ---------- 辅助函数 ----------
def get_heading_level(paragraph):
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        outline_lvl = pPr.find(qn('w:outlineLvl'))
        if outline_lvl is not None:
            return int(outline_lvl.get(qn('w:val'))) + 1
    style_name = paragraph.style.name.lower()
    for i in range(1, 10):
        if f'heading {i}' in style_name or f'标题 {i}' in style_name:
            return i
    return 0

def get_east_asian_font(run):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        return None
    rFonts = rPr.find(qn('w:rFonts'))
    return rFonts.get(qn('w:eastAsia')) if rFonts is not None else None

def para_in_table(paragraph):
    return paragraph._element.getparent().tag == qn('w:tc')

def section_has_even_and_odd_headers(section):
    return section._sectPr.find(qn('w:evenAndOddHeaders')) is not None

def section_has_title_pg(section):
    return section._sectPr.find(qn('w:titlePg')) is not None

def set_first_line_indent_xml(para, points):
    pPr = para._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.insert(0, ind)
    for attr in ['firstLine', 'firstLineChars', 'hanging', 'hangingChars']:
        if ind.get(qn(f'w:{attr}')) is not None:
            del ind.attrib[qn(f'w:{attr}')]
    ind.set(qn('w:firstLine'), str(int(points * 20)))

# ---------- 检查函数 ----------
def check_document(doc_path):
    errors = []
    try:
        doc = Document(doc_path)
    except Exception as e:
        return [("无法打开文档：" + str(e), False)]

    for i, section in enumerate(doc.sections):
        if abs(section.page_width.cm - 21.0) > 0.1 or abs(section.page_height.cm - 29.7) > 0.1:
            errors.append((f"第{i+1}节纸张大小不是A4", True))
        if (abs(section.top_margin.cm - 2.5) > 0.1 or
            abs(section.bottom_margin.cm - 2.5) > 0.1 or
            abs(section.left_margin.cm - 2.5) > 0.1 or
            abs(section.right_margin.cm - 2.5) > 0.1):
            errors.append((f"第{i+1}节页边距不符合2.5cm", True))

    background = doc.element.find(qn('w:background'))
    if background is not None:
        color = background.get(qn('w:color'), 'white')
        if color.lower() not in ('ffffff', 'auto'):
            errors.append(("文档背景颜色不是白色", True))

    for i, section in enumerate(doc.sections):
        if section_has_title_pg(section):
            h = section.first_page_header
            if h and any(p.runs for p in h.paragraphs):
                errors.append((f"第{i+1}节存在首页页眉", True))
        if section_has_even_and_odd_headers(section):
            eh = section.even_header
            oh = section.header
            if eh and any(p.runs for p in eh.paragraphs):
                errors.append((f"第{i+1}节存在偶数页页眉", True))
            if oh and any(p.runs for p in oh.paragraphs):
                errors.append((f"第{i+1}节存在奇数页页眉", True))
        else:
            h = section.header
            if h and any(p.runs for p in h.paragraphs):
                errors.append((f"第{i+1}节存在页眉", True))

        if section_has_title_pg(section):
            f = section.first_page_footer
            if f and any(p.runs for p in f.paragraphs):
                errors.append((f"第{i+1}节存在首页页脚", True))
        if section_has_even_and_odd_headers(section):
            ef = section.even_footer
            of = section.footer
            if ef and any(p.runs for p in ef.paragraphs):
                errors.append((f"第{i+1}节存在偶数页页脚", True))
            if of and any(p.runs for p in of.paragraphs):
                errors.append((f"第{i+1}节存在奇数页页脚", True))
        else:
            f = section.footer
            if f and any(p.runs for p in f.paragraphs):
                errors.append((f"第{i+1}节存在页脚", True))

    pb_count = 0
    for elem in doc.element.iter():
        if elem.tag == qn('w:br') and elem.get(qn('w:type')) == 'page':
            pb_count += 1
        elif elem.tag == qn('w:pageBreakBefore'):
            pb_count += 1
    if pb_count > 0:
        errors.append((f"文档包含{pb_count}处手动分页符", True))

    if len(doc.sections) > 1:
        errors.append(("文档存在分节符，不符合要求", False))

    for idx, para in enumerate(doc.paragraphs):
        if not para.text.strip() and not para.runs:
            continue
        pf = para.paragraph_format
        if pf.line_spacing_rule != WD_LINE_SPACING.EXACTLY:
            errors.append((f"段落{idx+1} 行距不是固定值", True))
        elif pf.line_spacing and abs(pf.line_spacing.pt - 30) > 0.5:
            errors.append((f"段落{idx+1} 行距应为30磅，实际{pf.line_spacing.pt}磅", True))
        if pf.space_before and pf.space_before.pt > 0:
            errors.append((f"段落{idx+1} 段前间距不为0", True))
        if pf.space_after and pf.space_after.pt > 0:
            errors.append((f"段落{idx+1} 段后间距不为0", True))
        if pf.alignment != WD_ALIGN_PARAGRAPH.LEFT:
            errors.append((f"段落{idx+1} 未左对齐", True))

        in_table = para_in_table(para)
        heading_level = get_heading_level(para) if not in_table else 0
        if heading_level == 1:
            expected_size = 16
            expected_indent = Pt(32)
        elif heading_level >= 2:
            expected_size = 14
            expected_indent = Pt(28)
        else:
            if in_table:
                expected_size = 12
                expected_indent = Pt(24)
            else:
                expected_size = 14
                expected_indent = Pt(28)

        actual_indent = pf.first_line_indent
        if actual_indent is None:
            errors.append((f"段落{idx+1} 未设置首行缩进", True))
        elif abs(actual_indent.pt - expected_indent.pt) > 1.0:
            errors.append((f"段落{idx+1} 首行缩进应为{expected_indent.pt}磅，实际{actual_indent.pt:.1f}磅", True))

        if ' ' in para.text:
            errors.append((f"段落{idx+1} 包含空格", False))

        if heading_level > 0:
            text = para.text.strip()
            if heading_level == 1:
                pattern = r'^\d+、'
            elif heading_level == 2:
                pattern = r'^\d+\.\d+、'
            elif heading_level == 3:
                pattern = r'^\d+\.\d+\.\d+、'
            else:
                pattern = None
            if pattern and not re.match(pattern, text):
                errors.append((f"标题段落{idx+1} 编号格式不正确", False))

        for r_idx, run in enumerate(para.runs):
            if not run.text.strip():
                continue
            east = get_east_asian_font(run)
            if east is not None and east != '宋体':
                errors.append((f"段落{idx+1} 运行{r_idx+1} 东亚字体非宋体({east})", not in_table))
            if run.font.name is not None and run.font.name != '宋体':
                errors.append((f"段落{idx+1} 运行{r_idx+1} 西文字体非宋体", not in_table))
            if run.font.size and abs(run.font.size.pt - expected_size) > 0.5:
                errors.append((f"段落{idx+1} 运行{r_idx+1} 字号应为{expected_size}磅", not in_table))
            if run.font.bold:
                errors.append((f"段落{idx+1} 运行{r_idx+1} 加粗", not in_table))
            if run.font.italic:
                errors.append((f"段落{idx+1} 运行{r_idx+1} 倾斜", not in_table))
            if run.font.underline:
                errors.append((f"段落{idx+1} 运行{r_idx+1} 下划线", not in_table))
            if run.font.color.rgb and run.font.color.rgb != RGBColor(0, 0, 0):
                errors.append((f"段落{idx+1} 运行{r_idx+1} 文字颜色非黑色", not in_table))
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None and rPr.find(qn('w:spacing')) is not None:
                errors.append((f"段落{idx+1} 运行{r_idx+1} 自定义字符间距", not in_table))

    for t_idx, table in enumerate(doc.tables):
        for c_idx, cell in enumerate(table._cells):
            tc = cell._tc
            borders = tc.find(qn('w:tcPr'))
            has_issue = False
            if borders is not None:
                for edge in ['top', 'left', 'bottom', 'right']:
                    b = borders.find(qn(f'w:{edge}'))
                    if b is None or b.get(qn('w:sz')) != '4' or b.get(qn('w:val')) != 'single':
                        has_issue = True
                        break
            else:
                has_issue = True
            if has_issue:
                errors.append((f"表格{t_idx+1} 单元格{c_idx+1} 边框不是0.5磅单线", True))

    errors.append(("【重要】无法自动计算文档页数，请手动确认总页数不超过500页，否则修复后仍不符合要求！", False))
    return errors

# ---------- 修复函数 ----------
def fix_document(input_path, output_path):
    doc = Document(input_path)

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    background = doc.element.find(qn('w:background'))
    if background is not None:
        doc.element.remove(background)

    for section in doc.sections:
        if section_has_title_pg(section):
            for part in [section.first_page_header, section.first_page_footer]:
                if part: part._element.clear()
        if section_has_even_and_odd_headers(section):
            for part in [section.even_header, section.even_footer, section.header, section.footer]:
                if part: part._element.clear()
        else:
            for part in [section.header, section.footer]:
                if part: part._element.clear()
        sectPr = section._sectPr
        for tag in ['w:titlePg', 'w:evenAndOddHeaders']:
            elem = sectPr.find(qn(tag))
            if elem is not None:
                sectPr.remove(elem)

    for elem in doc.element.iter():
        if elem.tag == qn('w:br') and elem.get(qn('w:type')) == 'page':
            elem.getparent().remove(elem)
    for pPr in doc.element.iter(qn('w:pPr')):
        pb = pPr.find(qn('w:pageBreakBefore'))
        if pb is not None:
            pPr.remove(pb)

    for para in doc.paragraphs:
        pf = para.paragraph_format
        in_table = para_in_table(para)

        pf.line_spacing = Pt(30)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

        heading_level = get_heading_level(para) if not in_table else 0
        if heading_level == 1:
            expected_size = 16
            target_indent = 32
        elif heading_level >= 2:
            expected_size = 14
            target_indent = 28
        else:
            if in_table:
                expected_size = 12
                target_indent = 24
            else:
                expected_size = 14
                target_indent = 28

        set_first_line_indent_xml(para, target_indent)

        if not in_table:
            for run in para.runs:
                if not run.text.strip():
                    continue
                rPr = run._element.find(qn('w:rPr'))
                need_fix = False
                east = get_east_asian_font(run)
                if east is not None and east != '宋体': need_fix = True
                if run.font.name is not None and run.font.name != '宋体': need_fix = True
                if run.font.size and abs(run.font.size.pt - expected_size) > 0.5: need_fix = True
                if run.font.bold: need_fix = True
                if run.font.italic: need_fix = True
                if run.font.underline: need_fix = True
                if run.font.color.rgb and run.font.color.rgb != RGBColor(0,0,0): need_fix = True
                if rPr is not None and rPr.find(qn('w:spacing')) is not None: need_fix = True

                if need_fix:
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        run._element.insert(0, rPr)
                    rFonts = rPr.find(qn('w:rFonts'))
                    if east is not None and east != '宋体':
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.insert(0, rFonts)
                        rFonts.set(qn('w:eastAsia'), '宋体')
                    if run.font.name is not None and run.font.name != '宋体':
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.insert(0, rFonts)
                        rFonts.set(qn('w:ascii'), '宋体')
                        rFonts.set(qn('w:hAnsi'), '宋体')
                        run.font.name = '宋体'
                    if run.font.size and abs(run.font.size.pt - expected_size) > 0.5:
                        run.font.size = Pt(expected_size)
                    if run.font.bold: run.font.bold = False
                    if run.font.italic: run.font.italic = False
                    if run.font.underline: run.font.underline = False
                    if run.font.color.rgb and run.font.color.rgb != RGBColor(0,0,0):
                        run.font.color.rgb = RGBColor(0,0,0)
                    for sp in rPr.findall(qn('w:spacing')):
                        rPr.remove(sp)

    for para in doc.paragraphs:
        in_table = para_in_table(para)
        heading_level = get_heading_level(para) if not in_table else 0
        if heading_level == 1:
            target = 32
        elif heading_level >= 2:
            target = 28
        else:
            target = 24 if in_table else 28
        set_first_line_indent_xml(para, target)

    for table in doc.tables:
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tblPr)
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        for edge in ['top', 'left', 'bottom', 'right']:
            edge_elem = tblBorders.find(qn(f'w:{edge}'))
            if edge_elem is None:
                edge_elem = OxmlElement(f'w:{edge}')
                tblBorders.append(edge_elem)
            edge_elem.set(qn('w:val'), 'single')
            edge_elem.set(qn('w:sz'), '4')
            edge_elem.set(qn('w:color'), '000000')

        for cell in table._cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            for edge in ['top', 'left', 'bottom', 'right']:
                b = tcBorders.find(qn(f'w:{edge}'))
                if b is None:
                    b = OxmlElement(f'w:{edge}')
                    tcBorders.append(b)
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')
                b.set(qn('w:color'), '000000')

    doc.save(output_path)

# ---------- 图形界面（检查项分类） ----------
class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Word格式检查与修复工具")
        self.root.geometry("900x700")
        self.file_path = tk.StringVar()

        frame_top = tk.Frame(root)
        frame_top.pack(pady=10, padx=10, fill=tk.X)
        tk.Label(frame_top, text="文档路径：").pack(side=tk.LEFT)
        tk.Entry(frame_top, textvariable=self.file_path, width=60).pack(side=tk.LEFT, padx=5)
        tk.Button(frame_top, text="选择文件", command=self.select_file).pack(side=tk.LEFT)
        tk.Button(frame_top, text="开始检查", command=self.run_check).pack(side=tk.LEFT, padx=5)

        self.text_area = scrolledtext.ScrolledText(root, wrap=tk.WORD, font=("宋体", 11))
        self.text_area.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        frame_bottom = tk.Frame(root)
        frame_bottom.pack(pady=10, fill=tk.X, padx=10)
        self.fix_button = tk.Button(frame_bottom, text="修复并另存为", command=self.fix_and_save, state=tk.DISABLED)
        self.fix_button.pack(side=tk.RIGHT, padx=5)
        tk.Button(frame_bottom, text="退出", command=root.quit).pack(side=tk.RIGHT)
        self.errors = []

    def select_file(self):
        path = filedialog.askopenfilename(filetypes=[("Word文档", "*.docx")])
        if path:
            self.file_path.set(path)

    def run_check(self):
        path = self.file_path.get()
        if not path or not os.path.exists(path):
            messagebox.showerror("错误", "文件不存在")
            return
        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(tk.END, "正在检查...\n")
        self.root.update()
        self.errors = check_document(path)
        self.text_area.delete(1.0, tk.END)

        if not self.errors:
            self.text_area.insert(tk.END, "✅ 文档完全符合规范。")
            self.fix_button.config(state=tk.DISABLED)
            return

        fixable_errors = [e for e in self.errors if e[1]]
        manual_errors = [e for e in self.errors if not e[1]]

        total = len(self.errors)
        self.text_area.insert(tk.END, f"共发现 {total} 处问题：\n\n")

        if fixable_errors:
            self.text_area.insert(tk.END, "---------- 可自动修复的问题 ----------\n")
            for i, (desc, _) in enumerate(fixable_errors, 1):
                self.text_area.insert(tk.END, f"{i}. [可修复] {desc}\n")
            self.text_area.insert(tk.END, f"\n自动修复项共 {len(fixable_errors)} 条。\n\n")

        if manual_errors:
            self.text_area.insert(tk.END, "---------- 检查项 ----------\n")
            for i, (desc, _) in enumerate(manual_errors, 1):
                self.text_area.insert(tk.END, f"{i}. [检查项] {desc}\n")
            self.text_area.insert(tk.END, f"\n检查项共 {len(manual_errors)} 条。\n")

        self.fix_button.config(state=tk.NORMAL if fixable_errors else tk.DISABLED)

    def fix_and_save(self):
        path = self.file_path.get()
        if not path:
            return

        if not messagebox.askyesno("页数确认",
            "重要提示：该文档的总页数是否已确认不超过500页？\n\n"
            "如果超过500页，修复后将不符合要求。\n"
            "请确保页数符合规范后再继续。"):
            return

        default_name = os.path.splitext(os.path.basename(path))[0] + "_修正版.docx"
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx")],
            initialdir=os.path.dirname(path),
            initialfile=default_name,
            title="保存修复后的文档"
        )
        if not save_path:
            return

        if os.path.abspath(save_path) == os.path.abspath(path):
            messagebox.showerror("错误", "不能覆盖原文件，请选择其他保存位置或修改文件名。")
            return

        try:
            self.text_area.insert(tk.END, "\n正在修复...\n")
            self.root.update()
            fix_document(path, save_path)
            self.text_area.insert(tk.END, f"修复完成，已保存至：{save_path}\n")
            messagebox.showinfo("完成", "修复成功，原文件未修改。")
        except Exception as e:
            messagebox.showerror("失败", str(e))

if __name__ == "__main__":
    root = tk.Tk()
    App(root)
    root.mainloop()